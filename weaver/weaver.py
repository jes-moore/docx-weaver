"""
Main Document Weaving Class
"""

from typing import Literal
import logging
from tqdm import tqdm
from docx import Document
from . import word
from .settings import DocxWeaverSettings
log = logging.getLogger(__name__)

class DocxWeaver:
    """
    Class to Convert/Translate and otherwise mutate Word Documents
    filename: str - Path to the Word Document
    purpose: str - Purpose of the Translation, used in all prompts
    paragraph_prompt: str | None - Prompt for Paragraphs (if any)
    table_prompt: str | None - Prompt for Tables (if any)
    mode: Literal["comments_only", "transform_only", "transform_and_comments"]
        - Mode of Operation
            - comments_only: Only add comments to the document
            - transform_only: Only transform the document inplace
            - transform_and_comments: Transform the document and put original text
            in comments
    """
    def __init__(
        self,
        filename: str,
        purpose: str,
        paragraph_prompt: str,
        table_prompt: str | None,
        mode: Literal["comments_only", "transform_only", "transform_and_comments"],
        openai_model_name: Literal["gpt-4-turbo", "gpt-3.5-turbo", "gpt-4o"] = "gpt-4o"
    ):
        assert mode in ["comments_only", "transform_only", "transform_and_comments"]
        assert isinstance(purpose, str)
        assert isinstance(paragraph_prompt, str)
        self.settings = DocxWeaverSettings(openai_model_name=openai_model_name)
        self.filename = filename
        self.document = Document(filename)
        self.table_prompt = table_prompt
        self.paragraph_prompt = paragraph_prompt
        self.purpose = purpose
        self.mode = mode

    def weave_document(self, output_fn: str):
        """
        Transforms the entire document
        """
        assert output_fn.endswith(".docx")

        # Paragraphs are always translated
        para_data = self._weave_paragraphs()
        section_para_data = self._weave_section_paragraphs()

        # Tables and Section Paragraphs/Headers are only run if transforming
        if self.mode in ["transform_only", "transform_and_comments"]:
            table_row_data = self._weave_tables()
            section_header_data = self._weave_section_headers()
        else:
            table_row_data = {}
            section_para_data = {}
            section_header_data = {}
        self.document.save(output_fn)
        log.info("Finished Weaving Document: %s", output_fn)
        # output_fn_unzipped = word.unpack_word_document(output_fn=output_fn)
        # word.rebuild_word_doc_from_zip(
        #     output_fn=output_fn,
        #     output_fn_unzipped=output_fn_unzipped
        # )
        return {
            "output_fn": output_fn,
            "paragraphs": para_data,
            "tables": table_row_data,
            "section_paragraphs": section_para_data,
            "section_headers": section_header_data
        }


    def _weave_paragraphs(self) -> dict[str, dict]:
        """
        Weave a single paragraph according to the prompt/constructor
        """
        # Translating Paragraphs
        log.info("Processing Paragraphs")
        paragraph_data = {}
        for ix_para, paragraph in tqdm(
            enumerate(self.document.paragraphs),
            total=len(self.document.paragraphs)
        ):
            if paragraph.text in ["", "\xa0", "\n"]:
                log.debug("No Processing For Paragraph = %s", ix_para)
                continue
            else:
                # Process and Insert Paragraph
                paragraph_data[str(ix_para)] = {
                    'type': "paragraph",
                    "runs": word.transform_paragraph(
                        paragraph=paragraph,
                        paragraph_prompt=self.paragraph_prompt,
                        purpose=self.purpose,
                        model_name=self.settings.openai_model_name,
                        mode=self.mode
                    )
                }
        log.info("Finished Processing Paragraphs")
        return paragraph_data

    def _weave_tables(self) -> dict[str, dict]:
        """
        Translate tables according to the prompt/construcot
        """
        log.info("Processing Tables")
        table_data = {}
        for ix_table, table in tqdm(
            enumerate(self.document.tables),
            total=len(self.document.tables)
        ):
            # Append Table
            table_data[str(ix_table)] = {
                "rows":  word.transform_table(
                    table,
                    table_prompt=self.table_prompt,
                    purpose=self.purpose,
                    model_name=self.settings.openai_model_name,
                    write_comments=True if "comments" in self.mode else False,
                )
            }
            log.debug("Finished Processing Table = %s\n", ix_table)
        log.info("Finished Processing Tables")
        return table_data

    def _weave_section_paragraphs(self) -> dict[str, dict]:
        """
        Convert/Transform all section paragraphs in the document
        """
        section_data = {}
        for ix_section, section in tqdm(
            enumerate(self.document.sections),
            total=len(self.document.sections)
        ):
            # Translate Headers
            header_paragraph_data = {}
            for ix_para, paragraph in enumerate(section.header.paragraphs):
                if "::::" in paragraph.text:
                    continue
                header_paragraph_data[str(ix_para)] = {
                    "type": "paragraph",
                    "runs": word.transform_paragraph(
                        paragraph,
                        paragraph_prompt=self.paragraph_prompt,
                        purpose=self.purpose,
                        model_name=self.settings.openai_model_name,
                        mode=self.mode,
                        root_type="header"
                    )
                }
            # Translate Footers
            footer_paragraph_data = {}
            for ix_para, paragraph in enumerate(section.footer.paragraphs):
                if "::::" in paragraph.text:
                    continue
                footer_paragraph_data[str(ix_para)] = {
                    "type": "paragraph",
                    "runs": word.transform_paragraph(
                        paragraph,
                        paragraph_prompt=self.paragraph_prompt,
                        purpose=self.purpose,
                        model_name=self.settings.openai_model_name,
                        mode=self.mode,
                        root_type="header"
                    )
                }
            # Translate First Page Header/Footer?
            first_page_header_paragraph_data = {}
            for ix_para, paragraph in enumerate(section.first_page_header.paragraphs):
                if "::::" in paragraph.text:
                    continue
                first_page_header_paragraph_data[str(ix_para)] = {
                    "type": "paragraph",
                    "runs": word.transform_paragraph(
                        paragraph,
                        paragraph_prompt=self.paragraph_prompt,
                        purpose=self.purpose,
                        model_name=self.settings.openai_model_name,
                        mode=self.mode,
                        root_type="header"
                    )
                }
            # Translate Footers
            first_page_footer_paragraph_data = {}
            for ix_para, paragraph in enumerate(section.first_page_footer.paragraphs):
                if "::::" in paragraph.text:
                    continue
                first_page_footer_paragraph_data[str(ix_para)] = {
                    "type": "paragraph",
                    "runs": word.transform_paragraph(
                        paragraph,
                        paragraph_prompt=self.paragraph_prompt,
                        purpose=self.purpose,
                        model_name=self.settings.openai_model_name,
                        mode=self.mode,
                        root_type="header"
                    )
                }
            # Append Translation Data
            section_data[str(ix_section)] = {
                "type": "section",
                "header_paragraphs": header_paragraph_data,
                "footer_paragraphs": footer_paragraph_data,
                "first_page_header_paragraphs": first_page_header_paragraph_data,
                "first_page_footer_paragraphs": first_page_footer_paragraph_data
            }
        log.info("Finished Processing Section Paragraphs")
        return section_data

    def _weave_section_headers(self) -> dict[str, dict]:
        """
        Convert/Transform all section headers in the document
        """
        section_data = {}

        for ix_section, section in tqdm(
            enumerate(self.document.sections),
            total=len(self.document.sections)
        ):
            # Translate Headers
            header_table_data = {}
            for ix_table, table in enumerate(section.header.tables):
                header_table_data[str(ix_table)] = {
                    "type": "table",
                    "runs": word.transform_table(
                        table,
                        table_prompt=self.table_prompt,
                        purpose=self.purpose,
                        model_name=self.settings.openai_model_name,
                        write_comments=True if "comments" in self.mode else False,
                        root_type="header"
                    )
                }
            # Translate Footers
            footer_table_data = {}
            for ix_table, table in enumerate(section.footer.tables):
                footer_table_data[str(ix_table)] = {
                    "type": "table",
                    "runs": word.transform_table(
                        table,
                        table_prompt=self.table_prompt,
                        purpose=self.purpose,
                        model_name=self.settings.openai_model_name,
                        write_comments=True if "comments" in self.mode else False,
                        root_type="header"
                    )
                }
            # Translate First Page Header/Footer?
            first_page_header_table_data = {}
            for ix_table, table in enumerate(section.first_page_header.tables):
                first_page_header_table_data[str(ix_table)] = {
                    "type": "table",
                    "runs": word.transform_table(
                        table,
                        table_prompt=self.table_prompt,
                        purpose=self.purpose,
                        model_name=self.settings.openai_model_name,
                        write_comments=True if "comments" in self.mode else False,
                        root_type="header"
                    )
                }
            # Translate Footers
            first_page_footer_table_data = {}
            for ix_table, table in enumerate(section.first_page_footer.tables):
                first_page_footer_table_data[str(ix_table)] = {
                    "type": "table",
                    "runs": word.transform_table(
                        table,
                        table_prompt=self.table_prompt,
                        purpose=self.purpose,
                        model_name=self.settings.openai_model_name,
                        write_comments=True if "comments" in self.mode else False,
                        root_type="header"
                    )
                }
            # Append Translation Data
            section_data[str(ix_section)] = {
                "type": "section",
                "header_tables": header_table_data,
                "footer_tables": footer_table_data,
                "first_page_header_tables": first_page_header_table_data,
                "first_page_footer_tables": first_page_footer_table_data
            }
        log.info("Finished Processing Section Headers")
        return section_data
