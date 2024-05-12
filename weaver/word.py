"""
Module for all word integration related
functions
"""

# General Imports
from typing import Any, Literal
import logging
import os
import shutil
import time
import copy
import string
import json
import pandas as pd
import openai
import docx

# Logger
log = logging.getLogger(__name__)


def transform_table(
    table,
    table_prompt: str,
    purpose: str,
    model_name: str,
    write_comments: bool,
    root_type: str = "table",
) -> dict[Any, Any]:
    """
    Primary function for translation a paragraph into
    the tgt language
    """
    row_data = {}
    transformed_texts = []  # Record For Merged/Duplicates
    for ix_row, row in enumerate(table.rows):
        row_cell_data = {}
        for ix_row_cell, cell in enumerate(row.cells):
            if cell.text.strip() in transformed_texts:
                # Skip This Strange Hook For Merged Cells
                continue
            # Aggregate transaltion across entire cell
            part_original = ""
            total_original = ""
            total_translation = ""
            row_cell_para_data = {}
            for ix_row_cell_para, paragraph in enumerate(cell.paragraphs):
                if "::::" in paragraph.text:
                    log.debug("Skipping Already Translated Paragraph")
                    continue
                if "Page" in paragraph.text:
                    log.debug("Skipping Already Translated Paragraph")
                    continue
                if paragraph.text in ["", "\xa0", "\n"]:
                    log.debug("No Processing For Input Paragraph")
                    continue
                if paragraph.text in transformed_texts:
                    continue
                else:
                    # Strip Mixed-Font Runs And Convert Runs Containing them.
                    cleanup_bad_runs(paragraph)

                    # Process Runs
                    row_cell_para_run_data = {}
                    for ix_row_cell_para_run, run in enumerate(paragraph.runs):
                        if "::::" in run.text:
                            log.debug("Skipping Already Translated Paragraph")
                            continue
                        if run.text.strip() in ["", "\xa0", ".", "$", "●"]:
                            continue
                        else:
                            # Store Comment Text For Later Update
                            original_text = copy.deepcopy(str(run.text))

                            # Transform Text
                            run.text, translated, add_comment = transform_text(
                                src_text=run.text,
                                prompt=table_prompt,
                                purpose=purpose,
                                model_name=model_name
                            )
                            run.text += f" :::: {run.text} ::::"
                            if translated:  # Record For Comment
                                part_original += original_text
                            total_original += original_text
                            total_translation += copy.deepcopy(str(run.text))
                            transformed_texts.append(total_translation.strip())
                        # Append Nested Run Data
                        row_cell_para_run_data[str(ix_row_cell_para_run)] = {
                            "original": original_text,
                            "translation": run.text,
                            "translated": translated
                        }
                    # Append Paragraph
                    row_cell_para_data[str(ix_row_cell_para)] = {
                        "runs": row_cell_para_run_data
                    }
            # Append Cell
            row_cell_data[str(ix_row_cell)] = {
                "paragraphs": row_cell_para_data
            }

            # Record Last Translation
            if total_translation.strip() != "":
                transformed_texts.append(total_translation)

            # Add Short Run Containing Comment
            if part_original != "":
                if len(cell.paragraphs) == 0:
                    pass
                else:
                    # Append Run, And Comment To Avoid Overwriting Formatting
                    last_para = cell.paragraphs[-1]
                    last_para.append_runs("")
                    run = last_para.runs[-1]
                    if add_comment:
                        if root_type not in ["header", "footer"]:
                            if write_comments:
                                run.add_comment(
                                    text=total_original,
                                    author='WordWeaver',
                                    initials="WW",
                                    dtime=pd.Timestamp.now(tz="UTC").strftime("%Y-%m-%d")
                                )
        # Append Row
        row_data[str(ix_row)] = {
            "cells": row_cell_data
        }
    return row_data


def transform_paragraph(
    paragraph: docx.text.paragraph.Paragraph,
    paragraph_prompt: str,
    purpose: str,
    model_name: str,
    mode: Literal["comments_only", "transform_only", "transform_and_comments"],
    root_type: str = "paragraph",
) -> None:
    """
    Primary function for translation a paragraph into
    the tgt language
    """

    # Cleanup Paragraph In Place
    cleanup_bad_runs(paragraph)

    # Process Runs
    run_data = {}
    for ix_run, run in enumerate(paragraph.runs):

        # Review What Changed Before Jan9, 2020 to try to keep linking.
        if run.text.strip() in ["", "\xa0", ".", "$", "●"]:
            continue
        else:
            # Store Comment Text For Later Update
            original_text = str(copy.deepcopy(run.text))

            if mode in ["comments_only"]:
                # Get Translation Only (Comment In this case)
                comment, translated, _ = transform_text(
                    src_text=run.text,
                    prompt=paragraph_prompt,
                    purpose=purpose,
                    model_name=model_name
                )
            else:
                # Update To Translate Text
                run.text, translated, _ = transform_text(
                    src_text=run.text,
                    prompt=paragraph_prompt,
                    purpose=purpose,
                    model_name=model_name
                )
                comment = original_text
            if translated:  # Record For Comment
                # Can't Add Comment To Header // Footer
                if root_type not in ["header", "footer"]:
                    if mode in ["transform_and_comments", "comments_only"]:
                        run.add_comment(
                            text=comment,
                            author='WordWeaver',
                            initials="WW",
                            dtime=pd.Timestamp.now(tz="UTC").strftime("%Y-%m-%d")
                        )
                else:
                    run.text += f" :::: {original_text} ::::"

            run_data[str(ix_run)] = {
                "original": original_text,
                "translation": run.text,
                "translated": translated,
            }


def cleanup_bad_runs(para):
    """
    Some runs get broken by italics/runs that make weaving
    not possible. This needs better handling, but this hack helps
    for now.
    """
    def check_country_abbr(para, jx) -> bool:
        """
        Check if sentence ends mid-country abbreviation (U. CA., etc)
        """
        try:
            is_country_abbr = (para.runs[jx-1].text.strip().endswith('.') &
                               para.runs[jx-1].text.strip()[-2].isupper())
        except Exception: # pylint: disable=broad-except
            is_country_abbr = False
        return is_country_abbr

    def check_special_end_cases(para, jx) -> bool:
        """
        Check special case with . then ,
        """
        try:
            is_special_case = (para.runs[jx-1].text.strip().endswith('.') &
                               para.runs[jx].text.strip().startswith(','))
        except Exception: # pylint: disable=broad-except
            is_special_case = False
        return is_special_case

    # Parse In Reverse
    for jx, _ in reversed(list(enumerate(para.runs))):
        # Stop At Start
        if jx == 0:
            break
        # If Previous Sentence Does Not End With A Period
        try:
            if (
                (not para.runs[jx - 1].text.strip().endswith("."))
                | check_country_abbr(para, jx)
                | check_special_end_cases(para, jx)
            ) & (para.runs[jx - 1].text not in ["\n", "\t"]):
                # Not Tab Or Newline
                if (not para.runs[jx].text.startswith("\t")) & (
                    not para.runs[jx].text.startswith("\n")
                ):
                    if (para.runs[jx].text not in ["\n", "\t"]) & (
                        para.runs[jx - 1].text not in ["\n", "\t"]
                    ):
                        if (jx - 1 == 0) and (para.runs[jx - 1].text == ""):
                            # Skip Case, No Use In Joining
                            continue
                        # Not Bracket Number
                        if not all(
                            [
                                para.runs[jx].text[0] == "(",
                                para.runs[jx].text[-1] == ")",
                            ]
                        ):
                            # Check Case Ending With U.S.
                            if not all(
                                [
                                    para.runs[jx - 1].text.strip().endswith("U.S."),
                                    para.runs[jx].text.strip()[0].isupper(),
                                ]
                            ):
                                # Append To Previous
                                para.runs[jx - 1].text = (
                                    para.runs[jx - 1].text + para.runs[jx].text
                                )
                                para._p.remove(para.runs[jx]._r)  # Delete Run pylint: disable=protected-access
        except IndexError:
            continue

    return


def transform_text(
    src_text: str,
    prompt: str,
    purpose: str,
    model_name: str
) -> tuple[str, bool, bool]:
    """
    This functions runs the translation of an unput run/text. It still needs
    some refactoring but this is a bit better...
    """

    # Try To Parse Cell Values // Check Formats Not Requiring Translation
    if check_formats_not_to_translate(copy.deepcopy(src_text)):
        return src_text, False, False

    # Translate
    src_text, transforms_dict = parse_and_prepare_src_text_transforms(src_text=src_text)
    tgt_text = generate_transformation(
        src_text=src_text,
        prompt=prompt,
        purpose=purpose,
        model_name=model_name
    )
    tgt_text = reapply_src_text_transforms(
        tgt_text=tgt_text,
        transforms_dict=transforms_dict
    )

    return tgt_text, True, True


def generate_transformation(
    src_text: str,
    prompt: str,
    purpose: str,
    model_name: str
) -> str:
    """
    Generates Text For A Given Prompt
    """
    user_prompt = (
        "You are a tool used to apply user-specified transformations to text."
        "The user can specify its purpose, the prompt and the input text."
        "Please respond with a json of the form {'tgt_text': 'translated text'}."
        f"""
        Task Purpose: {purpose}
        Prompt: {prompt}
        Input Text: {src_text}
        """
    )
    for i in range(5):
        try:
            completions = openai.chat.completions.create(
                    model=model_name,
                    messages=[{"role":"user","content":user_prompt}],
                    max_tokens=len(src_text) + 50,
                    n=1,
                    stop=None,
                    response_format={ "type": "json_object" },
            )
            message = completions.choices[0].message.content
            if message is None:
                raise ValueError("No Response From OpenAI")
            assert message is not None
            message = json.loads(message)
            assert isinstance(message, dict)
            if "tgt_text" not in message:
                raise ValueError("No Translation Found")
            break
        except Exception as e: # pylint: disable=broad-except
            if i < 5:
                time.sleep(1)  # wait for 1 second before trying again
                continue
            else:
                raise e  # if after 3 attempts it still fails, raise the exception.
    assert isinstance(message, dict)
    return message["tgt_text"]


def parse_and_prepare_src_text_transforms(src_text: str) -> tuple[str, dict[str, Any]]:
    """
    Parse and Prepare the Source Text for Translation
    """
    # Init
    transforms_dict: dict[str, Any] = {}

    # Count Special Cases and Remove Before Running (added back later)
    transforms_dict['ltab'] = src_text[0:1] == '\t'
    transforms_dict['rtab'] = src_text[-1:] == '\t'
    transforms_dict['lnewline'] = src_text[0:1] == '\n'
    transforms_dict['rnewline'] = src_text[-1:] == '\n'
    transforms_dict['lspace'] = \
        len(src_text) - len(src_text.lstrip()) - \
        transforms_dict['ltab'] - \
        transforms_dict['lnewline']
    transforms_dict['rspace'] = \
        len(src_text) - len(src_text.rstrip()) - \
        transforms_dict['rtab'] - \
        transforms_dict['rnewline']

    # Init Some Transforms
    transforms_dict['outer_quotes'] = False
    transforms_dict['outer_parens'] = False
    transforms_dict['outer_triangles'] = False
    transforms_dict['outer_square_parens'] = False
    transforms_dict['ltriangle'] = False
    transforms_dict['rcolon'] = False
    transforms_dict['rsemicolon'] = False
    transforms_dict['titled'] = False
    transforms_dict['appended_period'] = False
    transforms_dict['lpunct'] = transforms_dict['rpunct'] = None
    src_text = src_text.strip()

    # Find And Strip Outer Parens
    if (src_text[0] in ['“', '"']) & (src_text[-1] in ['”', '"']):
        transforms_dict['outer_quotes'] = True
        src_text = src_text[1:-1]

    if (src_text[0] == '(') & (src_text[-1] == ')'):
        transforms_dict['outer_parens'] = True
        src_text = src_text[1:-1]

    # Find And Strip Outer Triangles
    if (src_text[0] == '<') & (src_text[-1] == '>'):
        transforms_dict['outer_triangles'] = True
        src_text = src_text[1:-1]
    if (src_text[0] == '<') & (">" not in src_text):
        transforms_dict['ltriangle'] = True
        src_text = src_text[1:]

    # Find And Strip Outer Square Parens
    if (src_text[0] == '[') & (src_text[-1] == ']'):
        transforms_dict['outer_square_parens'] = True
        src_text = src_text[1:-1]

    # Check first and last identical punctuation
    (
        src_text,
        transforms_dict['lpunct'],
        transforms_dict['rpunct']
    ) = check_first_and_last_char_punct(src_text)

    if src_text[-1] == ':':
        transforms_dict['rcolon'] = True
        src_text = src_text[:-1]
    if src_text[-1] == ';':
        transforms_dict['rsemicolon'] = True
        src_text = src_text[:-1]

    # Append Period?
    if (
            (src_text[-1] not in string.punctuation) &
            (len(src_text.split()) > 3) & (not any([p in src_text for p in string.punctuation]))
       ):
        src_text += "."
        transforms_dict['appended_period'] = True

    # Check If Entire Input Is UpperCase
    if (src_text.isupper()) & ("US$" not in src_text):
        transforms_dict['titled'] = True
        src_text = src_text.lower()

    return src_text, transforms_dict


def check_first_and_last_char_punct(text: str) -> tuple[str, str | None, str | None]:
    """
    Check if the first and last characters of a string are
    identical punctuation marks, and return the string without
    these characters if they are.
    """
    # First check if first and last char are identical
    if text[0] == text[-1]:
        if (text[0] in string.punctuation) & (text[-1] in string.punctuation):
            lpunct = rpunct = text[0]
            text = text[1:-1]
            if text[0] == " ":
                lpunct += " "
                text = text[1:]
            if text[-1] == " ":
                rpunct = " " + rpunct
                text = text[:-1]
            return text, lpunct, rpunct
    return text, None, None


def check_formats_not_to_translate(src_text: str) -> bool:
    """
    Check if a string should not be translated based on
    certain rules
    """
    # Check only Numbers
    if all([c.isdigit() for c in src_text]):
        return True
    # If Only Company Name, Return as-is (Move this within Translation API?)
    elif check_only_company_name(input_str=copy.deepcopy(src_text)):
        return True
    # If No Letters, Return Exact Input
    elif check_no_letters_brackets(input_str=copy.deepcopy(src_text)):
        return True
    # If No Letters, Return Exact Input
    elif check_less_than_two_letters(input_str=copy.deepcopy(src_text)):
        return True
    else:
        return False


def check_only_company_name(input_str):
    """
    Check if an input string is a company name
    (and only a company name), and return it as-is.

    Intermediate string names need to be handled seperately
    """
    corp_abbr = ['Inc.', 'Corp.', 'Ltd.', 'llc']
    contains_abbr = any([abbr.upper() in input_str.upper().split() for abbr in corp_abbr])
    if (len(input_str.split()) <= 5) & (contains_abbr):
        return True
    else:
        return False


def check_no_letters_brackets(input_str):
    """
    Check if an input contains brackets, but no letters.
    These are generally tab cases and should be returned as-is.
    """
    test_str = "".join(input_str.replace("\n", "").replace("\t", "").split())
    contains_bracket = any([c in ['(', ')'] for c in test_str])
    contains_letters = any([c.isalpha() for c in test_str])

    if (contains_bracket) & (not contains_letters):
        return True
    else:
        return False


def check_less_than_two_letters(input_str):
    """
    Check if an input contains brackets, but no letters.
    These are generally tab cases and should be returned as-is.
    """
    test_str = "".join(input_str.replace("\n", "").replace("\t", "").split())
    sum_letters = sum([c.isalpha() for c in test_str])
    if sum_letters <= 1:
        return True
    else:
        return False
    
def reapply_src_text_transforms(tgt_text: str, transforms_dict: dict) -> str:
    """
    Re-append Adjustments made prior to transformation
    """
    if transforms_dict['appended_period'] & (tgt_text[-1] == "."):
        tgt_text = tgt_text[:-1]
    if transforms_dict['titled']:
        tgt_text = tgt_text.upper()
    if transforms_dict['rcolon']:
        tgt_text += ":"
    if transforms_dict['rsemicolon']:
        tgt_text += ";"
    if transforms_dict['ltriangle']:
        tgt_text = f"<{tgt_text}"
    if transforms_dict['lpunct'] is not None:
        tgt_text = transforms_dict['lpunct'] + tgt_text + transforms_dict['rpunct']
    if transforms_dict['outer_square_parens']:
        tgt_text = f"[{tgt_text}]"
    if transforms_dict['outer_parens']:
        tgt_text = f"({tgt_text})"
    if transforms_dict['outer_triangles']:
        tgt_text = f"<{tgt_text}>"
    if transforms_dict['outer_quotes']:
        tgt_text = f"“{tgt_text}”"
    if transforms_dict['rspace'] > 0:
        tgt_text += (' ' * transforms_dict['rspace'])
    if transforms_dict['lspace'] > 0:
        transforms_dict['tgt_text'] = (' ' * transforms_dict['lspace']) + tgt_text
    if transforms_dict['rnewline']:
        tgt_text += '\n'
    if transforms_dict['lnewline']:
        tgt_text = '\n' + tgt_text
    if transforms_dict['rtab']:
        tgt_text += '\t'
    if transforms_dict['ltab']:
        tgt_text = '\t' + tgt_text
    return tgt_text


def unpack_word_document(output_fn: str) -> str:
    """
    Unpacks Word Document and saves to zip
    directory for intermediate output
    """
    output_fn_unzipped = output_fn.replace(".docx", "")

    # Unzip Word Doc
    shutil.unpack_archive(
        filename=output_fn,
        extract_dir=output_fn_unzipped,
        format="zip")
    log.info("Unpacked Document")

    return output_fn_unzipped


def rebuild_word_doc_from_zip(output_fn: str, output_fn_unzipped: str):
    """
    Rebuilds Word Document From Zip Directory
    """
    # Unzip Word Doc
    shutil.make_archive(
        base_name=output_fn,
        root_dir=output_fn_unzipped,
        format="zip")

    # Rename to docx (remove append .zip)
    os.rename(
        src=f"{output_fn}.zip",
        dst=output_fn)

    # Remove Zip Directory
    shutil.rmtree(output_fn_unzipped)
    log.info("Original Document Updated")
